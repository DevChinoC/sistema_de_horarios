"""
ui/word/generador_word_docente.py
Genera un documento Word (.docx) con el horario semanal de un docente.

Estructura del documento:
  - Membrete (imagen de fondo ajustada a la página, opcional)
  - Encabezado: Plan de estudios, Semestre, Docente
  - Tabla semanal: Hora | Lunes | Martes | Miércoles | Jueves | Viernes
  - Celdas con color azul claro para materias ocupadas
"""

import os
from typing import Optional

from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from application.dto.horario_docente_dto import (
    HorarioDocenteFilaDTO, HorarioDocenteResumenDTO,
)

# ─────────────────────────────────────────────────────────────
# Constantes de estilo
# ─────────────────────────────────────────────────────────────
_AZUL_HDR = "3D5FD2"     # azul encabezado tabla
_AZUL_CELL = "B5CBF7"    # azul claro para celdas con materia
_BLANCO = "FFFFFF"
_NEGRO = "000000"

_DIAS = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes"]


def _format_hora(h24: str) -> str:
    """Convierte '08:00' → '8:00', '13:30' → '1:30'."""
    try:
        parts = h24.split(":")
        h, m = int(parts[0]), parts[1]
        if h == 0:
            return f"12:{m}"
        elif h <= 12:
            return f"{h}:{m}"
        else:
            return f"{h - 12}:{m}"
    except (ValueError, IndexError):
        return h24


def _set_cell_shading(cell, color: str) -> None:
    """Aplica color de fondo a una celda de tabla Word."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, bold: bool = False,
                   size: int = 8, color: str = _NEGRO,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """Escribe texto en una celda con formato."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Arial"

    # Reducir espaciado de párrafo
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)


# ─────────────────────────────────────────────────────────────
# Clase principal
# ─────────────────────────────────────────────────────────────
class GeneradorWordDocente:
    """Genera el documento Word con el horario semanal de un docente.

    Parámetros
    ----------
    resumen         : HorarioDocenteResumenDTO con datos del docente y filas
    ruta_membrete   : path a la imagen del membrete (puede ser None)
    ruta_salida     : path completo del .docx a generar
    """

    def __init__(
        self,
        resumen: HorarioDocenteResumenDTO,
        ruta_membrete: Optional[str],
        ruta_salida: str,
    ) -> None:
        self._resumen = resumen
        self._ruta_membrete = ruta_membrete
        self._ruta_salida = ruta_salida

    def generar(self) -> str:
        """Genera el documento Word y retorna la ruta del archivo creado."""
        doc = Document()

        # Configurar márgenes de página
        section = doc.sections[0]
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # ── Membrete (imagen de fondo) ────────────────────────
        tiene_membrete = bool(
            self._ruta_membrete and os.path.exists(self._ruta_membrete)
        )
        if tiene_membrete:
            # Agregar imagen del membrete como header
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = hp.add_run()
            run.add_picture(self._ruta_membrete, width=Inches(6.5))

            # Agregar espacio después del membrete
            doc.add_paragraph()

        # ── Encabezado de texto ────────────────────────────────
        r = self._resumen
        enc_style = doc.add_paragraph()
        enc_style.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = enc_style.add_run(f"Plan de estudios: {r.nombre_plan}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Arial"

        enc2 = doc.add_paragraph()
        enc2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = enc2.add_run(f"Semestre: {r.semestre}")
        run2.bold = True
        run2.font.size = Pt(11)
        run2.font.name = "Arial"

        enc3 = doc.add_paragraph()
        enc3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run3 = enc3.add_run(f"Docente: {r.nombre_docente}")
        run3.bold = True
        run3.font.size = Pt(11)
        run3.font.name = "Arial"

        doc.add_paragraph()  # Espacio

        # ── Construir tabla semanal ────────────────────────────
        self._agregar_tabla_horario(doc)

        # ── Guardar ───────────────────────────────────────────
        doc.save(self._ruta_salida)
        return self._ruta_salida

    def _agregar_tabla_horario(self, doc: Document) -> None:
        """Construye la tabla Hora | Lunes | … | Viernes."""
        filas = self._resumen.filas

        # Recopilar franjas horarias únicas, ordenadas
        franjas = sorted(set(
            (f.hora_inicio, f.hora_fin) for f in filas
            if f.hora_inicio and f.hora_fin
        ))

        if not franjas:
            p = doc.add_paragraph("Sin horarios registrados.")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return

        # Mapa rápido: (dia, hora_inicio) → nombre_materia
        mapa: dict[tuple, str] = {}
        for f in filas:
            if f.dia and f.hora_inicio:
                mapa[(f.dia, f.hora_inicio)] = f.nombre_materia

        n_filas = len(franjas) + 1  # +1 para header
        n_cols = len(_DIAS) + 1      # +1 para columna Hora

        tabla = doc.add_table(rows=n_filas, cols=n_cols)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.style = "Table Grid"

        # ── Header ─────────────────────────────────────────────
        header_labels = ["Hora"] + _DIAS
        for col_idx, label in enumerate(header_labels):
            cell = tabla.cell(0, col_idx)
            _set_cell_shading(cell, _AZUL_HDR)
            _set_cell_text(cell, label, bold=True, size=9,
                           color=_BLANCO)

        # ── Filas de datos ─────────────────────────────────────
        for row_idx, (hi, hf) in enumerate(franjas, start=1):
            # Columna Hora
            hora_txt = f"{_format_hora(hi)} - {_format_hora(hf)}"
            cell_hora = tabla.cell(row_idx, 0)
            _set_cell_text(cell_hora, hora_txt, bold=True, size=8)

            # Columnas por día
            for col_idx, dia in enumerate(_DIAS, start=1):
                materia = mapa.get((dia, hi), "")
                cell = tabla.cell(row_idx, col_idx)
                _set_cell_text(cell, materia, size=8,
                               alignment=WD_ALIGN_PARAGRAPH.LEFT)
                if materia:
                    _set_cell_shading(cell, _AZUL_CELL)

        # Ajustar ancho de columnas
        for row in tabla.rows:
            row.cells[0].width = Cm(2.5)
            for i in range(1, n_cols):
                row.cells[i].width = Cm(3.0)
